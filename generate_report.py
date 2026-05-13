"""
Generador de reporte Word (.docx) con resultados del análisis CPM/Crashing.

Lee el mismo archivo de Excel que `cpm_crashing.py`, ejecuta el análisis y
escribe un documento Word que presenta directamente los resultados
(duración, costo, ruta crítica, holguras, escenarios de costo óptimo y de
plazo extendido Tint + N).

Uso:
    python generate_report.py CrashingWorkshop.xlsx [salida.docx]
"""

from __future__ import annotations
import sys
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

from cpm_crashing import (
    compute_analysis,
    cpm_pass,
    fmt_num,
    path_duration,
    total_path_slack,
)


BODY_FONT = "Roboto Serif"
HIGHLIGHT_COLOR = RGBColor(0xC0, 0x50, 0x4D)
PLACEHOLDER_COLOR = RGBColor(0x1F, 0x4E, 0x79)

MESES_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def fecha_es(d: date) -> str:
    return f"Popayán, {d.day} de {MESES_ES[d.month - 1]} de {d.year}"


def set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color is not None:
        run.font.color.rgb = color


def add_centered(doc, text, size=12, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def add_text(doc, text, size=12, bold=False, italic=False, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_run_font(
        p.add_run(text), size=size, bold=bold, italic=italic, color=color,
    )
    return p


def add_section_heading(doc, text):
    add_text(doc, text, size=14, bold=True)


def add_blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def add_screenshot_placeholder(doc, descripcion):
    """Marcador simple para una captura de POM-QM."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run(f"[ Pegar aquí: {descripcion} ]"),
        size=11, bold=True, italic=True, color=PLACEHOLDER_COLOR,
    )
    add_blank(doc, 1)


def _preds_of(activities, successors):
    preds = {a: [] for a in activities}
    for s, succs in successors.items():
        for t in succs:
            preds.setdefault(t, []).append(s)
    return preds


def _pomqm_headers_rows(activities, preds_of, time_columns, time_dict_overrides=None):
    """Filas en el formato exacto de POM-QM (un predecesor por columna).

    time_columns: lista de (header, fn) donde fn(activity_name, info) → valor.
    """
    max_preds = max((len(p) for p in preds_of.values()), default=0)
    headers = ["Activity"] + [h for h, _ in time_columns]
    for i in range(max_preds):
        headers.append(f"Predecessor {i + 1}")

    rows = []
    for a in sorted(activities):
        info = activities[a]
        row = [a] + [fmt_num(fn(a, info)) for _, fn in time_columns]
        preds_sorted = sorted(preds_of.get(a, []))
        for i in range(max_preds):
            row.append(preds_sorted[i] if i < len(preds_sorted) else "")
        rows.append(row)
    return headers, rows


def _pomqm_setup_lines(doc, modulo, titulo_sugerido, n_actividades):
    add_text(doc, f"Ve a {modulo}.", bold=True, color=PLACEHOLDER_COLOR)
    add_text(doc, f"Título sugerido: «{titulo_sugerido}».", color=PLACEHOLDER_COLOR)
    add_text(
        doc, f"Number of Activities: {n_actividades}", color=PLACEHOLDER_COLOR,
    )
    add_text(
        doc,
        "Table Structure: Immediate predecessor list — Row Names: A, B, C, D, …",
        color=PLACEHOLDER_COLOR,
    )
    add_blank(doc)
    add_text(
        doc, "Datos a ingresar en POM-QM:", italic=True, color=PLACEHOLDER_COLOR,
    )


def add_pomqm_single_time(doc, data, titulo_sugerido, time_dict, captura_label):
    """Bloque Single Time Estimate con un diccionario de tiempos arbitrario."""
    activities = data["activities"]
    preds_of = _preds_of(activities, data["successors"])

    _pomqm_setup_lines(
        doc, "Single Time Estimate", titulo_sugerido, len(activities),
    )
    headers, rows = _pomqm_headers_rows(
        activities, preds_of,
        time_columns=[("Activity time", lambda a, _info: time_dict[a])],
    )
    add_table(doc, headers, rows, text_color=PLACEHOLDER_COLOR)
    add_blank(doc)
    add_screenshot_placeholder(
        doc, f"captura del data set en POM-QM ({captura_label})",
    )


def add_pomqm_setup_crash(doc, data, titulo_sugerido):
    """Bloque de configuración para Crashing (datos completos)."""
    activities = data["activities"]
    preds_of = _preds_of(activities, data["successors"])

    _pomqm_setup_lines(doc, "Crashing", titulo_sugerido, len(activities))
    headers, rows = _pomqm_headers_rows(
        activities, preds_of,
        time_columns=[
            ("Normal time", lambda _a, info: info["tiempo_normal"]),
            ("Crash time", lambda _a, info: info["tiempo_intensivo"]),
            ("Normal cost", lambda _a, info: info["costo_normal"]),
            ("Crash cost", lambda _a, info: info["costo_intensivo"]),
        ],
    )
    add_table(doc, headers, rows, text_color=PLACEHOLDER_COLOR)
    add_blank(doc)
    add_screenshot_placeholder(doc, "captura del data set de Crashing en POM-QM")


def add_table(doc, headers, rows, mark_rows=None, text_color=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(h), size=10, bold=True, color=text_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    mark_set = set(mark_rows or [])
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if (i - 1) in mark_set:
                color = HIGHLIGHT_COLOR
            else:
                color = text_color
            set_run_font(
                p.add_run(str(val)),
                size=10,
                bold=(i - 1) in mark_set,
                color=color,
            )
    return table


# ---------------------------------------------------------------------------
# Carátula
# ---------------------------------------------------------------------------

def write_cover(doc, titulo, autores, taller_n, curso, profesor):
    """Portada con el formato del Formato-Documento.docx institucional."""
    add_centered(doc, titulo.upper(), size=16, bold=True)
    add_blank(doc, 1)
    add_centered(doc, autores.upper(), size=16, bold=True)
    add_centered(
        doc,
        f"Taller número {taller_n} en el curso {curso.upper()}",
        size=16, bold=True,
    )
    add_blank(doc, 5)

    add_centered(doc, "Profesor:", size=12, bold=True)
    add_centered(doc, profesor.upper(), size=12, bold=True)
    add_blank(doc, 3)

    add_centered(doc, "Universidad del Cauca", size=12, bold=True)
    add_centered(doc, "Facultad de Ingeniería Electrónica y Telecomunicaciones", size=12, bold=True)
    add_centered(doc, "Departamento de Sistemas", size=12, bold=True)
    add_centered(doc, curso, size=12, bold=True)
    today = date.today()
    add_centered(
        doc,
        f"Popayán, {MESES_ES[today.month - 1]} {today.year}",
        size=12, bold=True,
    )

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Tablas
# ---------------------------------------------------------------------------

def write_input_table(doc, activities, successors):
    preds_of = {a: [] for a in activities}
    for s, succs in successors.items():
        for t in succs:
            preds_of.setdefault(t, []).append(s)

    headers = [
        "Actividad",
        "Predecesores",
        "T normal",
        "C normal miles $",
        "T intensivo",
        "C intensivo miles $",
    ]
    rows = []
    for a in sorted(activities):
        info = activities[a]
        preds = ", ".join(sorted(preds_of.get(a, []))) or "-"
        rows.append([
            a, preds,
            fmt_num(info["tiempo_normal"]),
            fmt_num(info["costo_normal"]),
            fmt_num(info["tiempo_intensivo"]),
            fmt_num(info["costo_intensivo"]),
        ])
    add_table(doc, headers, rows)


def write_paths_camino_table(doc, paths, time_dict):
    """Tabla de caminos: Camino | duración | holgura."""
    durations = [path_duration(p, time_dict) for p in paths]
    critical = max(durations)
    indexed = sorted(zip(durations, paths), key=lambda x: -x[0])

    headers = ["Camino", "Duración", "Holgura del camino"]
    rows = []
    mark_rows = []
    for i, (dur, p) in enumerate(indexed):
        slack = critical - dur
        if abs(slack) < 1e-9:
            mark_rows.append(i)
        rows.append(["-".join(p), fmt_num(dur), fmt_num(slack)])
    add_table(doc, headers, rows, mark_rows=mark_rows)
    return critical, indexed


def write_activity_slack_table(doc, activities, successors, time_dict):
    IC, TC, IL, TL, holgura, project_duration = cpm_pass(
        activities, successors, time_dict
    )
    headers = [
        "Actividad", "Duración",
        "Inicio Cercano", "Término Cercano",
        "Inicio Lejano", "Término Lejano",
        "Holgura", "¿Crítica?",
    ]
    rows = []
    mark_rows = []
    for i, a in enumerate(sorted(activities)):
        es_critica = abs(holgura[a]) < 1e-9
        if es_critica:
            mark_rows.append(i)
        rows.append([
            a,
            fmt_num(time_dict[a]),
            fmt_num(IC[a]), fmt_num(TC[a]),
            fmt_num(IL[a]), fmt_num(TL[a]),
            fmt_num(holgura[a]),
            "SÍ" if es_critica else "No",
        ])
    add_table(doc, headers, rows, mark_rows=mark_rows)
    return project_duration


def write_desintensification_table(doc, activities, sugerido, extensiones):
    headers = [
        "Actividad",
        "Tiempo sugerido",
        "Tiempo mínimo (crash)",
        "Costo de la actividad",
        "Días desintensificados",
        "Costo/día",
        "Ahorro",
    ]
    rows = []
    mark_rows = []
    for i, a in enumerate(sorted(activities)):
        info = activities[a]
        ahorro = info["costo_por_dia"] * extensiones[a]
        if extensiones[a] > 0:
            mark_rows.append(i)
        costo_act = info["costo_intensivo"] - ahorro
        rows.append([
            a,
            fmt_num(sugerido[a]),
            fmt_num(info["tiempo_intensivo"]),
            fmt_num(costo_act),
            fmt_num(extensiones[a]),
            fmt_num(info["costo_por_dia"]),
            fmt_num(ahorro),
        ])
    add_table(doc, headers, rows, mark_rows=mark_rows)


def critical_path_string(paths, time_dict):
    durations = [path_duration(p, time_dict) for p in paths]
    critical = max(durations)
    crit_paths = [p for d, p in zip(durations, paths) if abs(d - critical) < 1e-9]
    return "; ".join("-".join(p) for p in crit_paths), critical


# ---------------------------------------------------------------------------
# Secciones del cuerpo
# ---------------------------------------------------------------------------

def section_inputs(doc, data):
    add_section_heading(doc, "Datos de entrada del proyecto")
    add_text(
        doc,
        f"El proyecto consta de {len(data['activities'])} actividades. "
        "La siguiente tabla resume tiempos y costos en escenario normal e intensivo:",
    )
    write_input_table(doc, data["activities"], data["successors"])
    add_blank(doc)


def section_normal(doc, data, include_screenshots=True):
    crit_str, _ = critical_path_string(data["paths"], data["normal_time"])
    add_section_heading(doc, "Resultados con tiempos normales")
    add_text(
        doc,
        f"La duración del proyecto con tiempos normales es Tn = "
        f"{fmt_num(data['t_normal'])} días.",
    )
    add_text(
        doc,
        f"El costo total normal del proyecto es ${fmt_num(data['total_normal_cost'])}.",
    )
    add_text(doc, f"La(s) ruta(s) crítica(s) es(son): {crit_str}.")
    add_blank(doc)

    if include_screenshots:
        add_text(
            doc, "Soporte desde POM-QM (Single Time Estimate):",
            bold=True, color=PLACEHOLDER_COLOR,
        )
        add_pomqm_single_time(
            doc, data,
            titulo_sugerido="normal",
            time_dict=data["normal_time"],
            captura_label="tiempos normales",
        )
        add_text(
            doc,
            "Resolver con Solve y tomar las siguientes capturas:",
            color=PLACEHOLDER_COLOR,
        )
        add_screenshot_placeholder(
            doc,
            "Solution con Early Start, Early Finish, Late Start, Late Finish y "
            "Slack (actividades críticas en rojo)",
        )
        add_screenshot_placeholder(
            doc,
            "Precedence Graph con la ruta crítica resaltada",
        )
        add_screenshot_placeholder(
            doc,
            "Gantt Chart (Early Times) con tiempos normales",
        )
        add_blank(doc)

    add_text(doc, "Caminos del proyecto con sus holguras:", italic=True)
    write_paths_camino_table(doc, data["paths"], data["normal_time"])
    add_blank(doc)
    add_text(doc, "Holguras por actividad:", italic=True)
    write_activity_slack_table(
        doc, data["activities"], data["successors"], data["normal_time"]
    )
    doc.add_page_break()


def section_crash(doc, data, include_screenshots=True):
    crit_str_int, _ = critical_path_string(data["paths"], data["crash_time"])
    add_section_heading(doc, "Resultados con tiempos intensivos")
    add_text(
        doc,
        f"La duración mínima del proyecto al intensificar todas las actividades es "
        f"Tint = {fmt_num(data['t_optimal'])} días.",
    )
    add_text(
        doc,
        f"El costo intensivo máximo (todas las actividades al tiempo crash) es "
        f"${fmt_num(data['total_crash_cost'])}.",
    )
    add_text(
        doc,
        f"El costo óptimo para mantener Tint = {fmt_num(data['t_optimal'])} días es "
        f"${fmt_num(data['costo_optimo'])}.",
    )
    add_text(
        doc,
        f"La(s) ruta(s) crítica(s) con tiempos intensivos: {crit_str_int}.",
    )
    add_blank(doc)

    if include_screenshots:
        add_text(
            doc, "Soporte desde POM-QM (Crashing + Single Time Estimate):",
            bold=True, color=PLACEHOLDER_COLOR,
        )
        add_text(
            doc,
            "Paso 1. Configurar el módulo Crashing con todos los datos del "
            "proyecto (Normal time, Crash time, Normal cost y Crash cost). "
            "Esta es la fuente de los tiempos intensivos (columna «Crash "
            "time») que se usarán en el siguiente paso.",
            color=PLACEHOLDER_COLOR,
        )
        add_pomqm_setup_crash(doc, data, titulo_sugerido="crashing")

        add_text(
            doc,
            "Paso 2. Para visualizar la red con todas las actividades "
            "intensificadas (Tint a costo máximo), tomar la columna «Crash "
            "time» del paso anterior y usarla como «Activity time» en un "
            "nuevo Single Time Estimate:",
            color=PLACEHOLDER_COLOR,
        )
        add_pomqm_single_time(
            doc, data,
            titulo_sugerido="intensivo",
            time_dict=data["crash_time"],
            captura_label="tiempos intensivos",
        )
        add_text(
            doc,
            "Resolver con Solve y tomar las siguientes capturas:",
            color=PLACEHOLDER_COLOR,
        )
        add_screenshot_placeholder(
            doc,
            "Solution con tiempos intensivos (mostrar Slack y rutas críticas)",
        )
        add_screenshot_placeholder(
            doc,
            "Precedence Graph con la(s) ruta(s) crítica(s) para tiempos "
            "intensivos",
        )
        add_screenshot_placeholder(
            doc,
            "Gantt Chart con tiempos intensivos",
        )
        add_blank(doc)

    add_text(doc, "Caminos del proyecto con sus holguras:", italic=True)
    write_paths_camino_table(doc, data["paths"], data["crash_time"])
    add_blank(doc)
    add_text(doc, "Holguras por actividad:", italic=True)
    write_activity_slack_table(
        doc, data["activities"], data["successors"], data["crash_time"]
    )
    doc.add_page_break()


def section_optimal(doc, data, include_screenshots=True):
    crit_str_c, _ = critical_path_string(data["paths"], data["sugerido"])
    add_section_heading(
        doc,
        "Tabla sugerida de tiempos para el costo óptimo de Tint",
    )
    add_text(
        doc,
        f"Para alcanzar el costo óptimo de ${fmt_num(data['costo_optimo'])} "
        f"manteniendo Tint = {fmt_num(data['t_optimal'])} días, se sugieren "
        f"los siguientes tiempos por actividad. Las filas resaltadas son las "
        f"actividades desintensificadas (que se llevan a un tiempo mayor que el crash).",
    )
    write_desintensification_table(
        doc, data["activities"], data["sugerido"], data["extensiones"]
    )
    add_blank(doc)
    add_text(doc, f"Ahorro total = ${fmt_num(data['ahorros'])}.")
    diff = data["total_crash_cost"] - data["costo_optimo"]
    add_text(
        doc,
        f"Verificación: costo intensivo máximo − costo óptimo = "
        f"${fmt_num(data['total_crash_cost'])} − ${fmt_num(data['costo_optimo'])} "
        f"= ${fmt_num(diff)}.",
    )
    coincide = abs(data["ahorros"] - diff) < 1e-6
    add_text(
        doc,
        "La suma de ahorros individuales coincide con la diferencia entre el "
        "costo intensivo máximo y el costo óptimo, confirmando la solución."
        if coincide else
        "Nota: los valores no coinciden exactamente; revisar el modelo.",
        italic=True,
    )
    add_blank(doc)

    add_text(doc, f"Ruta(s) crítica(s) con los tiempos sugeridos: {crit_str_c}.")
    add_blank(doc)

    if include_screenshots:
        add_text(
            doc, "Soporte desde POM-QM (Crashing + Single Time Estimate):",
            bold=True, color=PLACEHOLDER_COLOR,
        )
        add_text(
            doc,
            "Paso 1. Usar el módulo Crashing ya configurado en la sección "
            "anterior. Pulsar Solve y luego ir a SOLUTIONS ▸ Crash Schedule. "
            "La columna «Crash by» indica cuántos días intensificar cada "
            f"actividad para alcanzar Tint = {fmt_num(data['t_optimal'])} "
            "días al mínimo costo. El «Crashing cost» total = sobrecosto "
            "óptimo respecto al costo normal (costo óptimo = "
            f"${fmt_num(data['costo_optimo'])}).",
            color=PLACEHOLDER_COLOR,
        )
        add_screenshot_placeholder(
            doc,
            f"Crash Schedule con «Crash by» y «Crashing cost» "
            f"(Project Crash time = {fmt_num(data['t_optimal'])})",
        )
        add_blank(doc)

        add_text(
            doc,
            "Paso 2. Para visualizar la red con los tiempos óptimos "
            "sugeridos, crear un nuevo Single Time Estimate usando como "
            "«Activity time» el tiempo sugerido de cada actividad (= Normal "
            "time − Crash by del paso 1):",
            color=PLACEHOLDER_COLOR,
        )
        add_pomqm_single_time(
            doc, data,
            titulo_sugerido="óptimo",
            time_dict=data["sugerido"],
            captura_label="tiempos óptimos para Tint",
        )
        add_text(
            doc,
            "Resolver con Solve y tomar las siguientes capturas:",
            color=PLACEHOLDER_COLOR,
        )
        add_screenshot_placeholder(
            doc,
            "Solution con los tiempos óptimos (rutas críticas resaltadas)",
        )
        add_screenshot_placeholder(
            doc,
            "Precedence Graph con la(s) ruta(s) crítica(s) en el escenario "
            "de costo óptimo",
        )
        add_screenshot_placeholder(
            doc,
            "Gantt Chart con los tiempos óptimos",
        )
        add_blank(doc)

    add_text(doc, "Caminos del proyecto con sus holguras:", italic=True)
    write_paths_camino_table(doc, data["paths"], data["sugerido"])
    add_blank(doc)
    add_text(doc, "Holguras por actividad:", italic=True)
    write_activity_slack_table(
        doc, data["activities"], data["successors"], data["sugerido"]
    )
    doc.add_page_break()


def section_extended_deadline(doc, data, include_screenshots=True):
    extra = data["extra_target_days"]
    crit_str_d, _ = critical_path_string(data["paths"], data["sugerido_target"])
    add_section_heading(
        doc,
        f"Resultados con plazo extendido (Tint + {extra} días)",
    )
    add_text(
        doc,
        f"Si se dispone de un plazo de Tint + {extra} = "
        f"{fmt_num(data['t_target'])} días, el costo óptimo del proyecto baja a "
        f"${fmt_num(data['costo_optimo_target'])}. Esto representa un ahorro "
        f"adicional de "
        f"${fmt_num(data['costo_optimo'] - data['costo_optimo_target'])} "
        f"respecto al costo óptimo para Tint.",
    )
    add_text(
        doc,
        f"La(s) ruta(s) crítica(s) con los nuevos tiempos: {crit_str_d}.",
    )
    add_blank(doc)

    add_text(
        doc,
        "Tabla sugerida de tiempos para el plazo extendido. Las filas resaltadas "
        "indican las actividades a desintensificar y cuántos días cada una "
        "respecto al escenario crash:",
        italic=True,
    )
    write_desintensification_table(
        doc, data["activities"], data["sugerido_target"], data["extensiones_target"]
    )
    add_blank(doc)

    if include_screenshots:
        add_text(
            doc, "Soporte desde POM-QM (Single Time Estimate):",
            bold=True, color=PLACEHOLDER_COLOR,
        )
        add_text(
            doc,
            "El módulo Crashing de POM-QM solo busca la duración mínima del "
            "proyecto, por lo que no puede usarse directamente para un plazo "
            f"objetivo distinto. La tabla de tiempos sugeridos para Tint + "
            f"{extra} = {fmt_num(data['t_target'])} días se obtiene por PL "
            "(ya incluida más arriba). Para visualizar la red con esos "
            "tiempos, crear un nuevo Single Time Estimate:",
            color=PLACEHOLDER_COLOR,
        )
        add_pomqm_single_time(
            doc, data,
            titulo_sugerido=f"plazo extendido (Tint + {extra})",
            time_dict=data["sugerido_target"],
            captura_label=f"tiempos para Tint + {extra}",
        )
        add_text(
            doc,
            "Resolver con Solve y tomar las siguientes capturas:",
            color=PLACEHOLDER_COLOR,
        )
        add_screenshot_placeholder(
            doc,
            "Solution con los tiempos del plazo extendido (rutas críticas "
            "resaltadas)",
        )
        add_screenshot_placeholder(
            doc,
            "Precedence Graph con la(s) ruta(s) crítica(s) en el plazo "
            "extendido",
        )
        add_screenshot_placeholder(
            doc,
            "Gantt Chart con los tiempos del plazo extendido",
        )
        add_blank(doc)

    add_text(doc, "Caminos del proyecto con sus holguras:", italic=True)
    write_paths_camino_table(doc, data["paths"], data["sugerido_target"])
    add_blank(doc)
    add_text(doc, "Holguras por actividad:", italic=True)
    write_activity_slack_table(
        doc, data["activities"], data["successors"], data["sugerido_target"]
    )
    doc.add_page_break()


def section_comparison(doc, data):
    extra = data["extra_target_days"]
    include_extended = extra > 0
    add_section_heading(doc, "Comparación de escenarios")
    if include_extended:
        add_text(
            doc,
            "La siguiente tabla resume los cuatro escenarios analizados: "
            "ejecución con tiempos normales, intensificación máxima, optimización "
            f"de costo para Tint, y optimización de costo para Tint + {extra} días.",
        )
    else:
        add_text(
            doc,
            "La siguiente tabla resume los tres escenarios analizados: "
            "ejecución con tiempos normales, intensificación máxima y "
            "optimización de costo para Tint.",
        )
    add_blank(doc)

    sobrecosto_b = data["total_crash_cost"] - data["total_normal_cost"]
    sobrecosto_c = data["costo_optimo"] - data["total_normal_cost"]
    dias_ahorrados = data["t_normal"] - data["t_optimal"]

    _, _, _, _, holg_a, _ = cpm_pass(
        data["activities"], data["successors"], data["normal_time"]
    )
    _, _, _, _, holg_b, _ = cpm_pass(
        data["activities"], data["successors"], data["crash_time"]
    )
    _, _, _, _, holg_c, _ = cpm_pass(
        data["activities"], data["successors"], data["sugerido"]
    )
    holg_total_cam_a, _ = total_path_slack(data["paths"], data["normal_time"])
    holg_total_cam_b, _ = total_path_slack(data["paths"], data["crash_time"])
    holg_total_cam_c, _ = total_path_slack(data["paths"], data["sugerido"])

    headers = [
        "Métrica",
        "Normal",
        "Tint costo máximo",
        "Tint costo óptimo",
    ]
    rows = [
        ["Duración del proyecto",
         fmt_num(data["t_normal"]),
         fmt_num(data["t_optimal"]),
         fmt_num(data["t_optimal"])],
        ["Costo total",
         f"${fmt_num(data['total_normal_cost'])}",
         f"${fmt_num(data['total_crash_cost'])}",
         f"${fmt_num(data['costo_optimo'])}"],
        ["Días ahorrados vs Tn",
         "0",
         fmt_num(dias_ahorrados),
         fmt_num(dias_ahorrados)],
        ["Sobrecosto vs costo normal",
         "$0",
         f"${fmt_num(sobrecosto_b)}",
         f"${fmt_num(sobrecosto_c)}"],
        ["Holgura total (actividades)",
         fmt_num(sum(holg_a.values())),
         fmt_num(sum(holg_b.values())),
         fmt_num(sum(holg_c.values()))],
        ["Margen de maniobra (suma holguras de caminos)",
         fmt_num(holg_total_cam_a),
         fmt_num(holg_total_cam_b),
         fmt_num(holg_total_cam_c)],
    ]

    if include_extended:
        sobrecosto_d = data["costo_optimo_target"] - data["total_normal_cost"]
        dias_ahorrados_d = data["t_normal"] - data["t_target"]
        _, _, _, _, holg_d, _ = cpm_pass(
            data["activities"], data["successors"], data["sugerido_target"]
        )
        holg_total_cam_d, _ = total_path_slack(
            data["paths"], data["sugerido_target"]
        )
        headers.append(f"Tint + {extra} costo óptimo")
        rows[0].append(fmt_num(data["t_target"]))
        rows[1].append(f"${fmt_num(data['costo_optimo_target'])}")
        rows[2].append(fmt_num(dias_ahorrados_d))
        rows[3].append(f"${fmt_num(sobrecosto_d)}")
        rows[4].append(fmt_num(sum(holg_d.values())))
        rows[5].append(fmt_num(holg_total_cam_d))
    else:
        holg_total_cam_d = None

    add_table(doc, headers, rows)
    add_blank(doc)

    add_text(doc, "Conclusiones del análisis:", bold=True)
    if sobrecosto_b > 0:
        ahorro_vs_b = sobrecosto_b - sobrecosto_c
        pct = 100 * ahorro_vs_b / sobrecosto_b if sobrecosto_b else 0
        add_text(
            doc,
            f"• Optimizar el costo manteniendo Tint = "
            f"{fmt_num(data['t_optimal'])} días genera un ahorro de "
            f"${fmt_num(ahorro_vs_b)} ({fmt_num(pct)}% menos sobrecosto) "
            f"respecto a intensificar todas las actividades.",
        )
    if include_extended and data["costo_optimo"] > data["costo_optimo_target"]:
        ahorro_extra = data["costo_optimo"] - data["costo_optimo_target"]
        add_text(
            doc,
            f"• Aceptar un plazo de {extra} días adicional "
            f"(Tint + {extra} = {fmt_num(data['t_target'])} días) reduce el "
            f"costo en ${fmt_num(ahorro_extra)} adicionales respecto al óptimo "
            f"de Tint.",
        )
    if holg_total_cam_b > holg_total_cam_c:
        add_text(
            doc,
            "• El escenario de intensificación máxima conserva más margen de "
            "maniobra (holgura total de caminos) que el escenario óptimo de "
            "Tint, pero a un costo significativamente mayor.",
        )
    if include_extended and holg_total_cam_d > holg_total_cam_c:
        add_text(
            doc,
            f"• El escenario con plazo extendido (Tint + {extra}) aumenta el "
            "margen de maniobra respecto al óptimo de Tint, lo que da mayor "
            "flexibilidad operativa además de menor costo.",
        )


def write_body(doc, data, include_screenshots=True):
    add_centered(doc, "Resultados del análisis", size=14, bold=True)
    add_blank(doc)

    section_inputs(doc, data)
    section_normal(doc, data, include_screenshots=include_screenshots)
    section_crash(doc, data, include_screenshots=include_screenshots)
    section_optimal(doc, data, include_screenshots=include_screenshots)
    if data.get("extra_target_days", 0) > 0:
        section_extended_deadline(
            doc, data, include_screenshots=include_screenshots
        )
    section_comparison(doc, data)


# ---------------------------------------------------------------------------
# Entrada principal
# ---------------------------------------------------------------------------

def build_report(
    xlsx_path,
    output_path,
    extra_target_days=5,
    titulo="Análisis de Crashing de Red CPM",
    autores="Nombres y Apellidos de los Estudiantes",
    taller_n="N",
    curso="Investigación de Operaciones",
    profesor="Ariel Pabón Burbano",
    include_screenshots=True,
):
    data = compute_analysis(xlsx_path, extra_target_days=extra_target_days)
    if data is None:
        raise RuntimeError("No se encontraron caminos en la red.")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(12)

    write_cover(doc, titulo, autores, taller_n, curso, profesor)
    write_body(doc, data, include_screenshots=include_screenshots)

    doc.save(output_path)
    return output_path


def run_interactive():
    """Modo interactivo: pregunta archivo, plazo, etc., al usuario."""
    import questionary

    here = Path(__file__).parent
    xlsx_files = sorted(
        p.name for p in here.glob("*.xlsx") if not p.name.startswith("~$")
    )
    if not xlsx_files:
        print("No se encontraron archivos .xlsx en", here)
        sys.exit(1)

    xlsx = questionary.select(
        "Selecciona el archivo Excel a analizar:",
        choices=xlsx_files,
    ).ask()
    if xlsx is None:
        sys.exit(0)

    default_out = f"Reporte_{Path(xlsx).stem}.docx"
    output = questionary.text(
        "Nombre del archivo de salida (.docx):",
        default=default_out,
    ).ask()
    if output is None:
        sys.exit(0)

    asignar_plazo = questionary.confirm(
        "¿Asignar un plazo objetivo Tint + N días (escenario d)?",
        default=True,
    ).ask()

    extra = 0
    if asignar_plazo:
        extra_str = questionary.text(
            "¿Cuántos días adicionales sobre Tint? (entero positivo)",
            default="5",
            validate=lambda v: v.isdigit() and int(v) >= 0
            or "Ingresa un número entero ≥ 0",
        ).ask()
        if extra_str is None:
            sys.exit(0)
        extra = int(extra_str)

    incluir_screenshots = questionary.confirm(
        "¿Incluir placeholders para capturas de POM-QM?",
        default=True,
    ).ask()

    titulo = questionary.text(
        "Título del trabajo (portada):",
        default="Análisis de Crashing de Red CPM",
    ).ask()
    autores = questionary.text(
        "Nombres y apellidos de los estudiantes (portada):",
        default="Nombres y Apellidos de los Estudiantes",
    ).ask()
    taller_n = questionary.text(
        "Número de taller (portada):",
        default="N",
    ).ask()
    curso = questionary.text(
        "Nombre del curso (portada):",
        default="Investigación de Operaciones",
    ).ask()
    profesor = questionary.text(
        "Profesor (portada):",
        default="Ariel Pabón Burbano",
    ).ask()

    print()
    print("Generando reporte con la siguiente configuración:")
    print(f"  Excel:           {xlsx}")
    print(f"  Salida:          {output}")
    print(f"  Plazo extendido: {'Tint + ' + str(extra) if extra else 'No (solo Tint)'}")
    print(f"  Screenshots:     {'Sí' if incluir_screenshots else 'No'}")
    print(f"  Título:          {titulo}")
    print(f"  Autores:         {autores}")
    print(f"  Taller número:   {taller_n}")
    print(f"  Curso:           {curso}")
    print(f"  Profesor:        {profesor}")
    print()

    generated = build_report(
        str(here / xlsx),
        str(here / output),
        extra_target_days=extra,
        titulo=titulo,
        autores=autores,
        taller_n=taller_n,
        curso=curso,
        profesor=profesor,
        include_screenshots=incluir_screenshots,
    )
    print(f"Reporte generado en: {generated}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        run_interactive()
    else:
        xlsx = sys.argv[1]
        out = sys.argv[2] if len(sys.argv) > 2 else "Reporte_Crashing.docx"
        extra = int(sys.argv[3]) if len(sys.argv) > 3 else 5
        generated = build_report(xlsx, out, extra_target_days=extra)
        print(f"Reporte generado en: {generated}")
